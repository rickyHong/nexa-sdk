from multiprocessing import Pool, cpu_count
from llama_index.core import SimpleDirectoryReader, Document
from llama_index.core.node_parser import TokenTextSplitter

import time
import subprocess
import fitz  # PyMuPDF
from PIL import Image
import pytesseract
import docx
from nexa.gguf import NexaVLMInference, NexaTextInference
import numpy as np
import json
import os
import re
import shutil

# Initialize the models
model_path = "llava-v1.6-vicuna-7b:q4_0"
model_path_text = "gemma-2b:q2_K"

inference = NexaVLMInference(
    model_path=model_path,
    local_path=None,
    stop_words=[],
    temperature=0.7,
    max_new_tokens=2048,
    top_k=50,
    top_p=1.0,
    profiling=True
)

inference_text = NexaTextInference(
    model_path=model_path_text,
    local_path=None,
    stop_words=[],
    temperature=0.7,
    max_new_tokens=512,
    top_k=50,
    top_p=0.9,
    profiling=True,
    embedding=True
)

def sanitize_description(description):
    # Remove invalid characters and limit the length of the file name
    sanitized = re.sub(r'[^\w\s-]', '', description).strip().lower()
    sanitized = re.sub(r'[-\s]+', '-', sanitized)
    return sanitized[:50]  # Limit to 50 characters

def get_response_text_from_generator(generator):
    response_text = ""
    try:
        while True:
            response = next(generator)
            choices = response.get('choices', [])
            for choice in choices:
                delta = choice.get('delta', {})
                if 'content' in delta:
                    response_text += delta['content']
    except StopIteration:
        pass
    return response_text

def read_word_file(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Dictionary to cache PDF content
pdf_cache = {}

def read_pdf_file(file_path):
    if file_path in pdf_cache:
        return pdf_cache[file_path]
    
    try:
        print(f"Attempting to read PDF file: {file_path}")
        doc = fitz.open(file_path)
        full_text = []
        for page in doc:
            full_text.append(page.get_text())
        pdf_content = '\n'.join(full_text)
        pdf_cache[file_path] = pdf_content
        print(f"Successfully read PDF file: {file_path}")
        return pdf_content
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
        return ""

def read_image_file(file_path):
    try:
        print(f"Attempting to read image file: {file_path}")
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image)
        print(f"Successfully read image file: {file_path}")
        return text
    except Exception as e:
        print(f"Error reading image file {file_path}: {e}")
        return ""

def read_text_file(file_path):
    with open(file_path, 'r') as file:
        text = file.read()
    return text

def process_document(args):
    file_path, chunk_size = args
    _, file_ext = os.path.splitext(file_path.lower())
    if file_ext == '.docx':
        text = read_word_file(file_path)
    elif file_ext == '.pdf':
        text = read_pdf_file(file_path)
    elif file_ext in ('.png', '.jpg', '.jpeg'):
        text = read_image_file(file_path)
    elif file_ext == '.txt':
        text = read_text_file(file_path)
    else:
        print(f"Skipping unsupported file type: {file_path}")
        return None, None
    
    if file_ext == '.pdf':
        # Do not split PDF files into chunks
        combined_text = text
    else:
        splitter = TokenTextSplitter(chunk_size=chunk_size)
        contents = splitter.split_text(text)
        combined_text = ' '.join(contents)
    
    return Document(text=combined_text, metadata={'file_path': file_path}), file_path

def load_documents_multiprocessing(path: str):
    if not os.path.exists(path):
        raise ValueError(f"Directory {path} does not exist.")
    
    # List all files in the directory
    file_paths = []
    for root, _, files in os.walk(path):
        for file in files:
            file_paths.append(os.path.join(root, file))
    
    # Limit to the first 10 files
    file_paths = file_paths[:10]
    
    reader = SimpleDirectoryReader(input_files=file_paths).iter_data()
    
    chunk_size = 6144
    with Pool(cpu_count()) as pool:
        results = pool.map(process_document, [(d.metadata['file_path'], chunk_size) for docs in reader for d in docs])
    
    # Filter out None results
    results = [result for result in results if result[0] is not None]
    
    documents = [document for document, _ in results]
    file_paths = [file_path for _, file_path in results]
    
    return documents, file_paths

def print_tree_with_subprocess(path):
    result = subprocess.run(['tree', path], capture_output=True, text=True)
    print(result.stdout)

def generate_image_description(image_path):
    description_generator = inference._chat(
        "Please provide a detailed description of this image in a sentence, emphasizing the meaning and context. Focus on capturing the key elements and underlying semantics.",
        image_path
    )
    description = get_response_text_from_generator(description_generator)
    return description

def get_decriptions_and_embeddings_for_images(image_paths):
    d = {}
    for image_path in image_paths:
        description = generate_image_description(image_path)
        embedding_result = inference_text.create_embedding(description)
        embedding = embedding_result["data"][0]['embedding']
        d[image_path] = {
            'description': description,
            'embedding': embedding
        }
    return d

# Recursive summarization function
MAX_CHUNK_SIZE = 2048  # Adjust based on your model's context window size
MAX_RECURSION_DEPTH = 5  # Prevent infinite recursion

def summarize_text_recursively(text, max_chunk_size=MAX_CHUNK_SIZE, recursion_depth=0):
    if recursion_depth > MAX_RECURSION_DEPTH:
        # Stop recursion and return text as is
        return text
    splitter = TokenTextSplitter(chunk_size=max_chunk_size)
    chunks = splitter.split_text(text)
    if len(chunks) == 1:
        # Text is short enough, generate summary directly
        summary = generate_summary(chunks[0])
        return summary
    else:
        summaries = []
        for chunk in chunks:
            summary = summarize_text_recursively(chunk, max_chunk_size, recursion_depth+1)
            summaries.append(summary)
        # Combine summaries and summarize again
        combined_summaries = ' '.join(summaries)
        # Now check if combined_summaries is short enough
        if len(splitter.split_text(combined_summaries)) <= 1:
            final_summary = generate_summary(combined_summaries)
            return final_summary
        else:
            # Continue recursion
            return summarize_text_recursively(combined_summaries, max_chunk_size, recursion_depth+1)

def generate_summary(text):
    description_generator = inference._chat(
        "Please provide a detailed summary of the following text in a sentence, emphasizing the key points and context.",
        text
    )
    description = get_response_text_from_generator(description_generator)
    return description

def generate_text_description(input_text):
    summary = summarize_text_recursively(input_text)
    return summary

def get_descriptions_and_embeddings_for_texts(text_tuples):
    results = []
    for file_path, text in text_tuples:
        description = generate_text_description(text)
        embeddings = inference_text.create_embedding(description)["data"][0]['embedding']
        results.append({
            'file_path': file_path,
            'description': description,
            'embeddings': embeddings
        })
    return results

def extract_topic(description):
    # Use VLMInference to extract the main topic from the description
    topic_generator = inference._chat(
        "Please provide the main topic of the following description in a single word or short phrase.",
        description
    )
    topic = get_response_text_from_generator(topic_generator)
    return topic.strip()

def create_directory_structure(base_path, topic):
    # Create a directory structure based on the topic
    dir_path = os.path.join(base_path, topic)
    os.makedirs(dir_path, exist_ok=True)
    return dir_path

if __name__ == '__main__':
    path = "/Users/q/nexa/nexa_sdk_local_file_organization/nexa-sdk/examples/local_file_organization/sample_data"
    new_path = "/Users/q/nexa/nexa_sdk_local_file_organization/nexa-sdk/examples/local_file_organization/renamed_files"
    
    if not os.path.exists(path):
        print(f"Directory {path} does not exist. Please create it and add the necessary files.")
    else:
        start_time = time.time()
        documents, file_paths = load_documents_multiprocessing(path)
        end_time = time.time()
        
        print(f"Time taken to load documents: {end_time - start_time:.2f} seconds")
        print("-"*50)
        print("Directory tree before renaming:")
        print_tree_with_subprocess(path)
        
        image_files = [doc.metadata['file_path'] for doc in documents if os.path.splitext(doc.metadata['file_path'].lower())[1] in ('.png', '.jpg', '.jpeg')]
        descriptions_and_embeddings_images = get_decriptions_and_embeddings_for_images(image_files)
        
        text_files = [doc.metadata['file_path'] for doc in documents if os.path.splitext(doc.metadata['file_path'].lower())[1] == '.txt']
        pdf_files = [doc.metadata['file_path'] for doc in documents if os.path.splitext(doc.metadata['file_path'].lower())[1] == '.pdf']
        
        # Create a list of tuples (file_path, text_content) for text and PDF files
        text_tuples = [(file_path, read_text_file(file_path)) for file_path in text_files]
        pdf_tuples = [(file_path, read_pdf_file(file_path)) for file_path in pdf_files]
        
        # Combine text and PDF tuples
        text_and_pdf_tuples = text_tuples + pdf_tuples
        
        descriptions_and_embeddings_texts = get_descriptions_and_embeddings_for_texts(text_and_pdf_tuples)
        
        output_file_images = "data/images_with_embeddings.json"
        os.makedirs(os.path.dirname(output_file_images), exist_ok=True)  # Ensure the directory exists
        with open(output_file_images, 'w') as f:
            json.dump(descriptions_and_embeddings_images, f, indent=4)
        
        output_file_texts = "data/texts_with_embeddings.json"
        os.makedirs(os.path.dirname(output_file_texts), exist_ok=True)  # Ensure the directory exists
        with open(output_file_texts, 'w') as f:
            json.dump(descriptions_and_embeddings_texts, f, indent=4)
        
        renamed_files = set()
        processed_files = set()  # Set to keep track of processed files
        
        # Ensure the new directory exists
        os.makedirs(new_path, exist_ok=True)
        
        # Counters for uploaded and generated files
        uploaded_files_count = len(file_paths)
        generated_files_count = 0
        
        for image_path, data in descriptions_and_embeddings_images.items():
            if image_path in processed_files:
                continue  # Skip if the file has already been processed
            processed_files.add(image_path)
            
            print(f"Image: {image_path}")
            print(f"Description: {data['description']}")
            # Extract topic from description
            topic = extract_topic(data['description'])
            # Create directory structure based on topic
            dir_path = create_directory_structure(new_path, topic)
            # Rename the file based on the description
            new_file_name = sanitize_description(data['description']) + os.path.splitext(image_path)[1]
            new_file_path = os.path.join(dir_path, new_file_name)
            
            # Ensure unique file name
            counter = 1
            while new_file_path in renamed_files or os.path.exists(new_file_path):
                new_file_name = f"{sanitize_description(data['description'])}_{counter}" + os.path.splitext(image_path)[1]
                new_file_path = os.path.join(dir_path, new_file_name)
                counter += 1
            
            shutil.copy2(image_path, new_file_path)
            renamed_files.add(new_file_path)
            generated_files_count += 1
            print(f"Copied and renamed to: {new_file_path}")
            print("-"*50)
        
        for text_data in descriptions_and_embeddings_texts:
            if text_data['file_path'] in processed_files:
                continue  # Skip if the file has already been processed
            processed_files.add(text_data['file_path'])
            
            print(f"File: {text_data['file_path']}")
            print(f"Description: {text_data['description']}")
            # Extract topic from description
            topic = extract_topic(text_data['description'])
            # Create directory structure based on topic
            dir_path = create_directory_structure(new_path, topic)
            # Rename the file based on the description
            new_file_name = sanitize_description(text_data['description']) + os.path.splitext(text_data['file_path'])[1]
            new_file_path = os.path.join(dir_path, new_file_name)
            
            # Ensure unique file name
            counter = 1
            while new_file_path in renamed_files or os.path.exists(new_file_path):
                new_file_name = f"{sanitize_description(text_data['description'])}_{counter}" + os.path.splitext(text_data['file_path'])[1]
                new_file_path = os.path.join(dir_path, new_file_name)
                counter += 1
            
            shutil.copy2(text_data['file_path'], new_file_path)
            renamed_files.add(new_file_path)
            generated_files_count += 1
            print(f"Copied and renamed to: {new_file_path}")
            print("-"*50)
        
        print("Directory tree after copying and renaming:")
        print_tree_with_subprocess(new_path)
        
        # Print the counts
        print(f"Number of files uploaded: {uploaded_files_count}")
        print(f"Number of files generated after renaming: {generated_files_count}")